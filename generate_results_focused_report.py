from __future__ import annotations

import argparse
from pathlib import Path
import re

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


TITLE = "Sonuc Odakli Benchmark Raporu"
SUBTITLE = "Yeni Jenness Turevi Yontemin Sonuclar Uzerinden Degerlendirilmesi"

METHOD_LABELS = {
    "gradient_multiplier": "Gradient Multiplier",
    "tin_2tri_cell": "TIN 2-Triangle",
    "jenness_window_8tri": "Klasik Jenness 8-Triangle",
    "sector_adaptive_jenness_integral": "Yeni Jenness Turevi (Sector-Adaptive)",
}

REALISTIC_PRESETS = {
    "mountain",
    "valley",
    "hills",
    "coastal",
    "plateau",
    "canyon",
    "volcanic",
    "glacial",
    "karst",
    "alluvial",
}

TEST_PRESETS = {
    "plane",
    "waves",
    "crater_field",
    "terraced",
    "patchwork",
    "mixed",
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate results-focused DOCX report.")
    parser.add_argument("--results-dir", default="sonuclar_out_", help="Directory containing batch_summary.xlsx")
    parser.add_argument("--output", default="", help="Output DOCX path")
    return parser.parse_args()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.9)
    section.bottom_margin = Cm(1.9)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    for style_name, size in (("Title", 20), ("Heading 1", 15), ("Heading 2", 12)):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(hdr[i], "D9EAF7")
        if widths is not None:
            hdr[i].width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
            if widths is not None:
                cells[i].width = Cm(widths[i])
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def format_pct(x: float, digits: int = 3) -> str:
    if pd.isna(x):
        return "-"
    return f"{x * 100:.{digits}f}%"


def format_sec(x: float) -> str:
    if pd.isna(x):
        return "-"
    if x >= 3600:
        return f"{x / 3600:.2f} saat"
    if x >= 60:
        return f"{x / 60:.2f} dk"
    return f"{x:.3f} sn"


def friendly_method(name: str) -> str:
    return METHOD_LABELS.get(name, name)


def extract_preset(area_id: str) -> str:
    text = str(area_id)
    if text.startswith("synth_"):
        text = text[len("synth_"):]
    return re.sub(r"_\d+x\d+_dx[^_]+_seed\d+_\d{8}_\d{6}$", "", text)


def friendly_preset(name: str) -> str:
    return name.replace("analytic_", "analytic ").replace("_", " ").title()


def family_for_preset(preset: str) -> str:
    if preset.startswith("analytic_"):
        return "Analitik"
    if preset in REALISTIC_PRESETS:
        return "Gercekci"
    if preset in TEST_PRESETS:
        return "Test Pattern"
    return "Diger"


def load_data(results_dir: Path) -> tuple[pd.DataFrame, pd.DataFrame]:
    batch = results_dir / "batch_summary.xlsx"
    long_df = pd.read_excel(batch, sheet_name="batch_results_long")
    summary_df = pd.read_excel(batch, sheet_name="batch_summary")
    return long_df, summary_df


def prepare_frames(long_df: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    nonref = long_df.loc[~long_df["is_reference_row"].fillna(False)].copy()
    nonref["preset"] = nonref["area_id"].map(extract_preset)
    nonref["family"] = nonref["preset"].map(family_for_preset)
    nonref["abs_rel_err"] = nonref["comparison_abs_rel_err"]
    comp = nonref.loc[nonref["abs_rel_err"].notna()].copy()
    return nonref, comp


def pairwise_metrics(comp: pd.DataFrame, method_a: str, method_b: str) -> dict[str, float]:
    pair = comp[comp["method"].isin([method_a, method_b])].pivot_table(
        index=["area_id", "gsd_m"],
        columns="method",
        values=["abs_rel_err", "runtime_sec"],
    )
    pair.columns = ["_".join(c) for c in pair.columns]
    pair = pair.dropna()
    a_err = pair[f"abs_rel_err_{method_a}"]
    b_err = pair[f"abs_rel_err_{method_b}"]
    runtime_ratio = pair[f"runtime_sec_{method_a}"] / pair[f"runtime_sec_{method_b}"]
    gain = (b_err - a_err) / b_err
    return {
        "comparisons": float(len(pair)),
        "strict_better": float((a_err < b_err).sum()),
        "better_1pct": float((gain > 0.01).sum()),
        "better_5pct": float((gain > 0.05).sum()),
        "worse_1pct": float((gain < -0.01).sum()),
        "median_runtime_ratio": float(runtime_ratio.median()),
        "median_gain": float(gain.median()),
        "mean_gain": float(gain.mean()),
    }


def build_sector_jenness_frame(comp: pd.DataFrame) -> pd.DataFrame:
    pair = comp[comp["method"].isin(["jenness_window_8tri", "sector_adaptive_jenness_integral"])].pivot_table(
        index=["area_id", "gsd_m", "preset", "family"],
        columns="method",
        values=["abs_rel_err", "runtime_sec"],
    )
    pair.columns = ["_".join(c) for c in pair.columns]
    pair = pair.dropna().reset_index()
    pair["err_diff"] = pair["abs_rel_err_sector_adaptive_jenness_integral"] - pair["abs_rel_err_jenness_window_8tri"]
    pair["runtime_ratio"] = pair["runtime_sec_sector_adaptive_jenness_integral"] / pair["runtime_sec_jenness_window_8tri"]
    pair["sector_better"] = pair["err_diff"] < 0
    pair["sector_gain"] = (
        pair["abs_rel_err_jenness_window_8tri"] - pair["abs_rel_err_sector_adaptive_jenness_integral"]
    ) / pair["abs_rel_err_jenness_window_8tri"]
    return pair


def generate_charts(results_dir: Path, summary_df: pd.DataFrame, sector_jenness: pd.DataFrame) -> dict[str, Path]:
    assets = results_dir / "report_assets_results"
    assets.mkdir(parents=True, exist_ok=True)
    chart_paths: dict[str, Path] = {}

    plt.style.use("seaborn-v0_8-whitegrid")

    native = summary_df.loc[summary_df["gsd_m"] == summary_df["gsd_m"].min()].copy()

    p = assets / "native_tradeoff.png"
    fig, ax = plt.subplots(figsize=(8.0, 5.0))
    ax.scatter(native["median_runtime_sec"], native["median_abs_rel_err"] * 100.0, s=110, color="#1f6fb2")
    for _, row in native.iterrows():
        ax.annotate(friendly_method(row["method"]), (row["median_runtime_sec"], row["median_abs_rel_err"] * 100.0), textcoords="offset points", xytext=(8, 6), fontsize=9)
    ax.set_xscale("log")
    ax.set_yscale("log")
    ax.set_xlabel("Native medyan runtime (sn)")
    ax.set_ylabel("Native medyan mutlak goreli hata (%)")
    ax.set_title("Native seviyede dogruluk-hiz dengesi")
    fig.tight_layout()
    fig.savefig(p, dpi=180)
    plt.close(fig)
    chart_paths["tradeoff"] = p

    native_pair = sector_jenness.loc[sector_jenness["gsd_m"] == sector_jenness["gsd_m"].min()].copy()
    native_pair = native_pair.sort_values("err_diff")

    p = assets / "sector_vs_classic_jenness_native_error_diff.png"
    fig, ax = plt.subplots(figsize=(10.2, 5.6))
    labels = [friendly_preset(v) for v in native_pair["preset"]]
    colors = ["#2f9e44" if v < 0 else "#d9480f" for v in native_pair["err_diff"]]
    ax.bar(range(len(native_pair)), native_pair["err_diff"] * 100.0, color=colors)
    ax.axhline(0.0, color="black", linewidth=1.0)
    ax.set_xticks(range(len(native_pair)))
    ax.set_xticklabels(labels, rotation=75, ha="right", fontsize=8)
    ax.set_ylabel("Sector - klasik Jenness hata farki (yuzde puan)")
    ax.set_title("Native seviyede yeni Jenness turevi ile klasik Jenness farki")
    fig.tight_layout()
    fig.savefig(p, dpi=180)
    plt.close(fig)
    chart_paths["native_diff"] = p

    p = assets / "sector_vs_classic_jenness_native_runtime_ratio.png"
    fig, ax = plt.subplots(figsize=(10.2, 5.6))
    ax.bar(range(len(native_pair)), native_pair["runtime_ratio"], color="#5c7cfa")
    ax.axhline(1.0, color="black", linewidth=1.0)
    ax.set_xticks(range(len(native_pair)))
    ax.set_xticklabels(labels, rotation=75, ha="right", fontsize=8)
    ax.set_yscale("log")
    ax.set_ylabel("Runtime orani (sector / klasik Jenness)")
    ax.set_title("Native seviyede yeni Jenness turevinin runtime maliyeti")
    fig.tight_layout()
    fig.savefig(p, dpi=180)
    plt.close(fig)
    chart_paths["native_runtime"] = p

    g2_rows: list[dict[str, float]] = []
    for gsd, sub in sector_jenness.groupby("gsd_m", sort=True):
        g2_rows.append(
            {
                "gsd_m": float(gsd),
                "wins": int(sub["sector_better"].sum()),
                "total": int(len(sub)),
                "median_ratio": float(np.median(sub["abs_rel_err_sector_adaptive_jenness_integral"] / sub["abs_rel_err_jenness_window_8tri"])),
                "median_runtime_ratio": float(np.median(sub["runtime_ratio"])),
            }
        )
    g2 = pd.DataFrame(g2_rows)

    p = assets / "sector_vs_classic_jenness_gsd.png"
    fig, axes = plt.subplots(1, 2, figsize=(11.0, 4.6))
    axes[0].bar(g2["gsd_m"].astype(str), g2["wins"], color="#2f9e44")
    axes[0].set_xlabel("GSD (m)")
    axes[0].set_ylabel("Sector'in daha iyi oldugu alan sayisi")
    axes[0].set_title("GSD bazinda kazanim sayisi")
    axes[1].plot(g2["gsd_m"], g2["median_ratio"], marker="o", color="#d9480f", label="Hata orani")
    axes[1].plot(g2["gsd_m"], g2["median_runtime_ratio"], marker="s", color="#1f6fb2", label="Runtime orani")
    axes[1].set_xscale("log")
    axes[1].set_yscale("log")
    axes[1].set_xlabel("GSD (m)")
    axes[1].set_ylabel("Sector / klasik Jenness")
    axes[1].set_title("Medyan hata ve runtime oranlari")
    axes[1].legend()
    fig.tight_layout()
    fig.savefig(p, dpi=180)
    plt.close(fig)
    chart_paths["gsd"] = p

    selected = ["patchwork", "mixed", "canyon", "terraced"]
    p = assets / "selected_cases_sector_vs_classic_jenness.png"
    fig, axes = plt.subplots(2, 2, figsize=(10.6, 7.2), sharex=True)
    for ax, preset in zip(axes.flat, selected):
        sub = sector_jenness.loc[sector_jenness["preset"] == preset].sort_values("gsd_m")
        ax.plot(sub["gsd_m"], sub["abs_rel_err_jenness_window_8tri"] * 100.0, marker="o", label="Klasik Jenness")
        ax.plot(sub["gsd_m"], sub["abs_rel_err_sector_adaptive_jenness_integral"] * 100.0, marker="s", label="Yeni Jenness turevi")
        ax.set_xscale("log")
        ax.set_yscale("log")
        ax.set_title(friendly_preset(preset))
        ax.set_xlabel("GSD (m)")
        ax.set_ylabel("Mutlak goreli hata (%)")
        ax.legend(fontsize=8)
    fig.tight_layout()
    fig.savefig(p, dpi=180)
    plt.close(fig)
    chart_paths["selected"] = p

    return chart_paths


def build_report(results_dir: Path, output: Path, nonref: pd.DataFrame, comp: pd.DataFrame, summary_df: pd.DataFrame, charts: dict[str, Path]) -> None:
    doc = Document()
    style_document(doc)

    p = doc.add_paragraph()
    p.style = doc.styles["Title"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(TITLE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(SUBTITLE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Tarih: 2026-03-23\n")
    p.add_run("Kaynak: sonuclar_out_/batch_summary.xlsx ve per-run sonuclar")

    doc.add_page_break()

    native_gsd = float(nonref["gsd_m"].min())
    native_summary = summary_df.loc[summary_df["gsd_m"] == native_gsd].sort_values("median_abs_rel_err").copy()
    sector_pair = build_sector_jenness_frame(comp)
    native_pair = sector_pair.loc[sector_pair["gsd_m"] == native_gsd].copy()

    pair_sector_vs_jenness = pairwise_metrics(comp, "sector_adaptive_jenness_integral", "jenness_window_8tri")
    pair_sector_vs_grad = pairwise_metrics(comp, "sector_adaptive_jenness_integral", "gradient_multiplier")
    pair_sector_vs_tin = pairwise_metrics(comp, "sector_adaptive_jenness_integral", "tin_2tri_cell")
    pair_jenness_vs_grad = pairwise_metrics(comp, "jenness_window_8tri", "gradient_multiplier")

    win_rows = []
    idx = comp.groupby(["area_id", "gsd_m"])["abs_rel_err"].idxmin()
    wins = comp.loc[idx].groupby("method").size().sort_values(ascending=False)
    for method, count in wins.items():
        win_rows.append([friendly_method(method), str(int(count))])

    doc.add_heading("1. Sonuc Ozeti", level=1)
    doc.add_paragraph(
        "Bu rapor yalnizca sayisal sonuclara odaklanir. Deney duzeni, kod mimarisi ve uygulama ayrintilari geri plana "
        "alinmis; bunun yerine benchmark ciktisinin bize ne soyledigi, ozellikle de yeni Jenness turevi yontemin pratik "
        "olarak ne kazandirdigi tartisilmistir."
    )
    add_bullets(
        doc,
        [
            f"Toplam degerlendirme alanlarinda genel lider {friendly_method('jenness_window_8tri')}: 88 karsilastirmanin 77'sinde en dusuk hata.",
            f"Yeni Jenness turevi yalnizca 88 karsilastirmanin 6'sinda klasik Jenness'ten daha iyi; bunlarin yalniz 2 tanesi anlamli ustunluk uretir.",
            f"Yeni Jenness turevi klasik Jenness'e gore medyan olarak yaklasik {pair_sector_vs_jenness['median_runtime_ratio']:.1f} kat daha yavas.",
            f"Yeni yontem, Gradient Multiplier'a karsi daha dogru olsa da bu kazanimin bedeli medyan {pair_sector_vs_grad['median_runtime_ratio']:.1f} kat ek runtime.",
            "Bu benchmark setinde yeni yontem genel-purpose bir replacement gibi degil, belirli zor karisik yuzeyler icin nis alan araci gibi davranir.",
        ],
    )

    doc.add_heading("2. Degerlendirme Kapsami", level=1)
    add_table(
        doc,
        headers=["Olcut", "Deger"],
        rows=[
            ["Senaryo sayisi", str(int(nonref["area_id"].nunique()))],
            ["Toplam aktif yontem", str(int(nonref["method"].nunique()))],
            ["Toplam GSD", str(int(nonref["gsd_m"].nunique()))],
            ["Referansli karsilastirma sayisi", str(int(comp[["area_id", "gsd_m"]].drop_duplicates().shape[0]))],
            ["Analitik yuzey sayisi", str(int((nonref["family"] == "Analitik").groupby(nonref["area_id"]).max().sum()))],
            ["Native GSD", f"{native_gsd:g} m"],
        ],
        widths=[6.0, 9.8],
    )

    doc.add_heading("3. Butun Yontemler Arasinda Genel Sonuc", level=1)
    add_table(
        doc,
        headers=["Yontem", "Native medyan hata", "Native ort. hata", "Native medyan runtime"],
        rows=[
            [
                friendly_method(row["method"]),
                format_pct(row["median_abs_rel_err"], 4),
                format_pct(row["mean_abs_rel_err"], 4),
                format_sec(row["median_runtime_sec"]),
            ]
            for _, row in native_summary.iterrows()
        ],
        widths=[5.4, 3.4, 3.4, 3.2],
    )
    doc.add_paragraph()
    add_table(doc, headers=["Alan bazli birincilik", "Kazanma sayisi"], rows=win_rows, widths=[8.6, 3.4])
    doc.add_paragraph()
    doc.add_picture(str(charts["tradeoff"]), width=Cm(15.6))
    add_caption(doc, "Sekil 1. Native seviyede tum yontemlerin dogruluk-hiz dengesi.")
    doc.add_paragraph(
        "Native seviyede tablo nettir: klasik Jenness hem en dusuk medyan hatayi verir hem de yeni Jenness turevinden "
        "cok daha dusuk runtime ile calisir. Yeni yontem, TIN ve Gradient'e gore daha sofistike olmasina ragmen bu "
        "benchmark setinde en iyi genel dengeyi kuran yontem klasik Jenness olmustur."
    )

    doc.add_heading("4. Yeni Jenness Turevi Yontemin Pairwise Degerlendirmesi", level=1)
    add_table(
        doc,
        headers=["Karsilastirma", "Toplam", "Daha iyi", ">%1 daha iyi", ">%5 daha iyi", ">%1 daha kotu", "Medyan runtime orani"],
        rows=[
            [
                "Yeni Jenness / Klasik Jenness",
                str(int(pair_sector_vs_jenness["comparisons"])),
                str(int(pair_sector_vs_jenness["strict_better"])),
                str(int(pair_sector_vs_jenness["better_1pct"])),
                str(int(pair_sector_vs_jenness["better_5pct"])),
                str(int(pair_sector_vs_jenness["worse_1pct"])),
                f"{pair_sector_vs_jenness['median_runtime_ratio']:.1f}x",
            ],
            [
                "Yeni Jenness / Gradient",
                str(int(pair_sector_vs_grad["comparisons"])),
                str(int(pair_sector_vs_grad["strict_better"])),
                str(int(pair_sector_vs_grad["better_1pct"])),
                str(int(pair_sector_vs_grad["better_5pct"])),
                str(int(pair_sector_vs_grad["worse_1pct"])),
                f"{pair_sector_vs_grad['median_runtime_ratio']:.1f}x",
            ],
            [
                "Yeni Jenness / TIN",
                str(int(pair_sector_vs_tin["comparisons"])),
                str(int(pair_sector_vs_tin["strict_better"])),
                str(int(pair_sector_vs_tin["better_1pct"])),
                str(int(pair_sector_vs_tin["better_5pct"])),
                str(int(pair_sector_vs_tin["worse_1pct"])),
                f"{pair_sector_vs_tin['median_runtime_ratio']:.1f}x",
            ],
            [
                "Klasik Jenness / Gradient",
                str(int(pair_jenness_vs_grad["comparisons"])),
                str(int(pair_jenness_vs_grad["strict_better"])),
                str(int(pair_jenness_vs_grad["better_1pct"])),
                str(int(pair_jenness_vs_grad["better_5pct"])),
                str(int(pair_jenness_vs_grad["worse_1pct"])),
                f"{pair_jenness_vs_grad['median_runtime_ratio']:.1f}x",
            ],
        ],
        widths=[5.0, 1.8, 1.6, 2.0, 2.0, 2.0, 2.2],
    )
    doc.add_paragraph()
    add_bullets(
        doc,
        [
            "Yeni Jenness turevi, Gradient Multiplier'a karsi beklenen sekilde daha dogru; ancak ayni dogruluk kazanci klasik Jenness ile de zaten elde ediliyor.",
            "Kritik soru yeni yontemin klasik Jenness'e karsi ne kattigi. Mevcut batch bu soruya zayif bir cevap veriyor: yalnizca 6/88 strict win ve bunlarin yalniz 2'si %5'in ustunde anlamli iyilesme.",
            "Buna karsin yeni yontem klasik Jenness'ten 21 durumda %1'den fazla daha kotu ve 14 durumda %5'ten fazla daha kotu.",
        ],
    )

    doc.add_heading("5. Yeni Jenness Turevi ve Klasik Jenness: Native Seviyede Dogrudan Fark", level=1)
    doc.add_picture(str(charts["native_diff"]), width=Cm(16.0))
    add_caption(doc, "Sekil 2. Native seviyede sector-adaptive Jenness ile klasik Jenness arasindaki hata farki. Negatif deger yeni yontemin daha iyi oldugunu gosterir.")
    doc.add_picture(str(charts["native_runtime"]), width=Cm(16.0))
    add_caption(doc, "Sekil 3. Native seviyede yeni Jenness turevinin klasik Jenness'e gore runtime maliyeti.")
    doc.add_paragraph(
        "Native seviyedeki resim daha da belirgin. Yeni yontem anlamli kazanci sadece Patchwork ve Mixed yuzeylerinde veriyor. "
        "Buna karsin Canyon ve Terraced gibi iki kritik zor yuzeyde klasik Jenness'in belirgin sekilde gerisine dusuyor. "
        "Bu, yeni yontemin genel 'zor yuzey cozucusu' olmadigini; belirli karisik ama yumusak gecisli kompozit yuzeylerde "
        "avantaj saglayip, keskin kirik geometri iceren yuzeylerde ise ters tepebildigini gosteriyor."
    )

    native_best = native_pair.sort_values("err_diff").head(6)
    native_worst = native_pair.sort_values("err_diff", ascending=False).head(6)
    add_table(
        doc,
        headers=["Yeni yontemin en iyi oldugu ornekler", "Hata farki", "Runtime orani"],
        rows=[
            [friendly_preset(row["preset"]), format_pct(row["err_diff"], 4), f"{row['runtime_ratio']:.1f}x"]
            for _, row in native_best.iterrows()
        ],
        widths=[7.4, 3.4, 3.4],
    )
    doc.add_paragraph()
    add_table(
        doc,
        headers=["Yeni yontemin en zayif kaldigi ornekler", "Hata farki", "Runtime orani"],
        rows=[
            [friendly_preset(row["preset"]), format_pct(row["err_diff"], 4), f"{row['runtime_ratio']:.1f}x"]
            for _, row in native_worst.iterrows()
        ],
        widths=[7.4, 3.4, 3.4],
    )

    doc.add_heading("6. GSD Bazinda Yeni Jenness Turevi", level=1)
    doc.add_picture(str(charts["gsd"]), width=Cm(16.0))
    add_caption(doc, "Sekil 4. GSD bazinda yeni Jenness turevinin klasik Jenness'e karsi kazanim sayisi ve oranlari.")
    doc.add_paragraph(
        "GSD buyudukce yeni yontemin klasik Jenness'e karsi pratik avantaji kayboluyor. 0.1-2 m araliginda gorulen az sayidaki ustunlukler "
        "esas olarak analitik ve neredeyse esit durumlar; farklar sayisal olarak yok denecek kadar kucuk. 5 m ve sonrasinda ise yeni yontem "
        "hicbir GSD seviyesinde klasik Jenness'i gecemiyor. Bu da yeni yaklasimin coarser-resolution rejiminde ek deger uretmedigini gosteriyor."
    )

    doc.add_heading("7. Secili Zor Ornekler", level=1)
    doc.add_picture(str(charts["selected"]), width=Cm(16.0))
    add_caption(doc, "Sekil 5. Dort kritik senaryoda yeni Jenness turevi ve klasik Jenness hata egirileri.")
    add_bullets(
        doc,
        [
            "Patchwork: yeni yontemin en anlamli kazanc sagladigi alan; klasik Jenness'e gore belirgin hata azalimi var.",
            "Mixed: patchwork benzeri sekilde yeni yontem burada da olumlu ayrisiyor.",
            "Canyon: yeni yontemin en sert basarisizliklarindan biri; klasik Jenness belirgin bicimde daha dogru.",
            "Terraced: keskin basamakli geometri yeni yontem icin problemli; hata klasik Jenness'in cok ustune cikiyor.",
        ],
    )

    doc.add_heading("8. Yuzey Ailesi Bazinda Yorum", level=1)
    family_rows = []
    for family in ["Analitik", "Gercekci", "Test Pattern"]:
        sub = native_pair.loc[native_pair["family"] == family]
        family_rows.append([
            family,
            str(int(sub["sector_better"].sum())),
            str(int(len(sub))),
            format_pct(sub["err_diff"].median(), 4),
            f"{sub['runtime_ratio'].median():.1f}x",
        ])
    add_table(
        doc,
        headers=["Aile", "Yeni yontem daha iyi", "Toplam", "Medyan hata farki", "Medyan runtime orani"],
        rows=family_rows,
        widths=[3.4, 3.1, 2.0, 3.6, 3.5],
    )
    doc.add_paragraph()
    add_bullets(
        doc,
        [
            "Analitik yuzeylerde iki yontem neredeyse esit. Buradaki sonuc, yeni yontemin puruzsuz benchmarklarda klasik Jenness'i asmaktan cok ona yaklastigini gosteriyor.",
            "Gercekci yuzey grubunda yeni yontem native seviyede hic kazanamiyor. Ustelik maliyeti dramatik sekilde daha yuksek.",
            "Test pattern grubunda tablo ikiye ayriliyor: patchwork ve mixed yararli, terraced ve crater/waves grubunda ise yarar sinirli ya da negatif.",
        ],
    )

    doc.add_heading("9. Son Hukum", level=1)
    doc.add_paragraph(
        "Mevcut benchmark ciktisina dayanarak yeni Jenness turevi yontem icin en dengeli yorum sudur: yontem arastirma acisindan "
        "ilgi cekici ve belirli karisik kompozit yuzeylerde anlamli kazanc uretebiliyor; ancak bu haliyle klasik Jenness'in genel "
        "amacli yerine gecen yeni varsayilan cozum oldugu soylenemez."
    )
    add_bullets(
        doc,
        [
            "Eger hedef tek bir varsayilan yontem secmekse, mevcut batch klasik Jenness'i destekliyor.",
            "Eger hedef zor karisik yuzeylerde secmeli bir ileri-yontem sunmaksa, yeni Jenness turevi bunu hakli cikaracak ilk isaretleri veriyor.",
            "Bildiride yeni yontemi 'replacement' yerine 'specialized extension' veya 'adaptive variant for mixed surfaces' olarak konumlandirmak daha guclu olabilir.",
            "Bir sonraki deney setinde patchwork/mixed benzeri yuzeyler cogaltilir ve canyon/terraced benzeri kirik geometri icin neden geride kaldigi incelenirse, makaledeki tartisma bolumu cok daha ikna edici olur.",
        ],
    )

    doc.save(output)


def main() -> None:
    args = parse_args()
    results_dir = Path(args.results_dir).resolve()
    output = Path(args.output) if args.output else results_dir / "sonuc_odakli_jenness_degerlendirme_raporu_2026-03-23.docx"
    output = output.resolve()

    long_df, summary_df = load_data(results_dir)
    nonref, comp = prepare_frames(long_df)
    charts = generate_charts(results_dir, summary_df, build_sector_jenness_frame(comp))
    build_report(results_dir, output, nonref, comp, summary_df, charts)
    print(f"Report generated: {output}")


if __name__ == "__main__":
    main()
