from __future__ import annotations

import argparse
import json
import re
from dataclasses import dataclass
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


REPORT_TITLE = "DEM/DSM 3B Yuzey Alani Hesaplama Calismasi Ara Raporu"
REPORT_SUBTITLE = "Kod tabani ve sonuclar_out_ ciktilari uzerinden uretilmistir"

METHOD_LABELS = {
    "gradient_multiplier": "Gradient Multiplier",
    "tin_2tri_cell": "TIN 2-Triangle",
    "jenness_window_8tri": "Jenness 8-Triangle",
    "sector_adaptive_jenness_integral": "Sector-Adaptive Jenness Integral",
    "bilinear_patch_integral": "Bilinear Patch Integral",
    "adaptive_bilinear_patch_integral": "Adaptive Bilinear Integral",
    "multiscale_decomposed_area": "Multiscale Decomposed Area",
}

METHOD_NOTES = {
    "gradient_multiplier": "Egim tabanli alan carpani; en hizli temel cizgi.",
    "tin_2tri_cell": "Her hucreyi iki ucgen olarak modelliyor.",
    "jenness_window_8tri": "3x3 pencerede 8 ucgen; mevcut batch'in en dengeli yontemi.",
    "sector_adaptive_jenness_integral": "3x3 quadratic fit ve sektor bazli adaptif integral.",
    "bilinear_patch_integral": "Bilinear patch uzerinde sayisal integral; kodda mevcut, bu batch'te yok.",
    "adaptive_bilinear_patch_integral": "Tolerans kontrollu adaptif bilinear integral; kodda mevcut, bu batch'te yok.",
    "multiscale_decomposed_area": "Topo + mikro ayristirma; ozel analizler icin mevcut, bu batch'te yok.",
}

MODULE_SUMMARY = [
    ("main.py", "Varsayilan config, method presetleri ve toplu calistirma giris noktasi."),
    ("surface_area/cli.py", "CLI orkestrasyonu, GSD dongusu, Excel/JSON/PNG ciktilari."),
    ("surface_area/methods.py", "Blok bazli raster akisi ve 6 temel alan algoritmasi."),
    ("surface_area/multiscale.py", "Topografik + mikro ayristirmasi."),
    ("surface_area/synthetic.py", "Gercekci/test/analitik sentetik DSM uretimi."),
    ("surface_area/analytic_surfaces.py", "Surekli ground truth verebilen analitik yuzeyler."),
    ("surface_area/plotting.py", "A3D, hata, runtime ve surface excess grafikleri."),
    ("surface_area/roi.py", "ROI/parsel maskesi ve secmeli alan raporlama akisi."),
    ("surface_area/io.py", "Raster okuma, yeniden ornekleme ve blok pencere islemleri."),
    ("tests/", "61 test fonksiyonundan olusan dogrulama dizisi."),
]

REALISTIC_PRESETS = {
    "mountain",
    "valley",
    "hills",
    "coastal",
    "plateau",
    "canyon",
    "volcanic",
    "glacial",
    "karst",
    "alluvial",
}

TEST_PRESETS = {
    "plane",
    "waves",
    "crater_field",
    "terraced",
    "patchwork",
    "mixed",
}


@dataclass(frozen=True)
class ReportStats:
    run_count: int
    method_count: int
    gsd_count: int
    non_reference_rows: int
    reference_rows: int
    analytic_count: int
    realistic_count: int
    test_pattern_count: int
    native_ref_count: int
    continuous_gt_count: int
    width: int
    height: int
    dx: float
    dy: float
    extent_width_m: float
    extent_height_m: float
    total_compute_sec: float
    total_unique_resample_sec: float


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate a Turkish DOCX progress report from sonuclar_out_ outputs.")
    parser.add_argument("--results-dir", default="sonuclar_out_", help="Results root containing batch_summary.xlsx.")
    parser.add_argument("--output", default="", help="Output DOCX path.")
    return parser.parse_args()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    for style_name, size in (("Title", 20), ("Heading 1", 15), ("Heading 2", 12)):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)


def add_title_page(doc: Document, output_name: str) -> None:
    title = doc.add_paragraph()
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(REPORT_TITLE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(REPORT_SUBTITLE)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Hazirlanma tarihi: 2026-03-23\n")
    meta.add_run(f"Belge: {output_name}\n")
    meta.add_run("Kaynak klasor: sonuclar_out_")

    doc.add_page_break()


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], col_widths_cm: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(hdr_cells[i], "D9EAF7")
        if col_widths_cm is not None:
            hdr_cells[i].width = Cm(col_widths_cm[i])

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            if col_widths_cm is not None:
                cells[i].width = Cm(col_widths_cm[i])

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def format_pct(value: float, digits: int = 3) -> str:
    if pd.isna(value):
        return "-"
    return f"{value * 100:.{digits}f}%"


def format_duration(seconds: float) -> str:
    if pd.isna(seconds):
        return "-"
    if seconds >= 3600:
        return f"{seconds / 3600:.2f} saat"
    if seconds >= 60:
        return f"{seconds / 60:.2f} dk"
    return f"{seconds:.2f} sn"


def friendly_method(name: str) -> str:
    return METHOD_LABELS.get(name, name)


def extract_preset(area_id: str) -> str:
    text = area_id
    if text.startswith("synth_"):
        text = text[len("synth_"):]
    return re.sub(r"_\d+x\d+_dx[^_]+_seed\d+_\d{8}_\d{6}$", "", text)


def family_for_preset(preset: str) -> str:
    if preset.startswith("analytic_"):
        return "Analitik"
    if preset in REALISTIC_PRESETS:
        return "Gercekci"
    if preset in TEST_PRESETS:
        return "Test pattern"
    return "Diger"


def friendly_preset(preset: str) -> str:
    return preset.replace("analytic_", "analytic ").replace("_", " ").title()


def load_run_infos(results_dir: Path, batch_runs: pd.DataFrame) -> list[dict[str, object]]:
    infos: list[dict[str, object]] = []
    for run_dir in batch_runs["run_dir"]:
        run_name = str(run_dir).split("\\")[-1]
        run_info_path = results_dir / run_name / "run_info.json"
        if not run_info_path.exists():
            continue
        infos.append(json.loads(run_info_path.read_text(encoding="utf-8")))
    return infos


def build_stats(
    long_df: pd.DataFrame,
    batch_runs: pd.DataFrame,
    run_infos: list[dict[str, object]],
) -> ReportStats:
    nonref = long_df.loc[~long_df["is_reference_row"].fillna(False)].copy()
    unique_presets = sorted({extract_preset(area_id) for area_id in nonref["area_id"].dropna().unique()})

    info0 = run_infos[0]
    dem_info = info0["dem_info"]
    width = int(dem_info["width"])
    height = int(dem_info["height"])
    dx = float(dem_info["dx"])
    dy = float(dem_info["dy"])

    uniq_resample = nonref.sort_values(["area_id", "gsd_m", "method"]).drop_duplicates(["area_id", "gsd_m"])

    return ReportStats(
        run_count=int(batch_runs["area_id"].nunique()),
        method_count=int(nonref["method"].nunique()),
        gsd_count=int(nonref["gsd_m"].nunique()),
        non_reference_rows=int(len(nonref)),
        reference_rows=int(long_df["is_reference_row"].fillna(False).sum()),
        analytic_count=sum(p.startswith("analytic_") for p in unique_presets),
        realistic_count=sum(p in REALISTIC_PRESETS for p in unique_presets),
        test_pattern_count=sum(p in TEST_PRESETS for p in unique_presets),
        native_ref_count=int(nonref.groupby("area_id")["synthetic_native_ref_A3D"].apply(lambda s: s.notna().any()).sum()),
        continuous_gt_count=int(nonref.groupby("area_id")["continuous_gt_A3D"].apply(lambda s: s.notna().any()).sum()),
        width=width,
        height=height,
        dx=dx,
        dy=dy,
        extent_width_m=width * dx,
        extent_height_m=height * dy,
        total_compute_sec=float(nonref["runtime_sec"].sum()),
        total_unique_resample_sec=float(uniq_resample["resample_runtime_sec"].fillna(0).sum()),
    )


def generate_charts(
    assets_dir: Path,
    batch_summary: pd.DataFrame,
    native_family_summary: pd.DataFrame,
) -> dict[str, Path]:
    assets_dir.mkdir(parents=True, exist_ok=True)

    error_chart = assets_dir / "batch_median_error_vs_gsd.png"
    runtime_chart = assets_dir / "batch_median_runtime_vs_gsd.png"
    tradeoff_chart = assets_dir / "native_accuracy_runtime_tradeoff.png"
    family_chart = assets_dir / "native_family_error.png"

    summary = batch_summary.copy()

    plt.style.use("seaborn-v0_8-whitegrid")

    fig, ax = plt.subplots(figsize=(8.8, 5.2))
    for method, g in summary.groupby("method", sort=False):
        ax.plot(g["gsd_m"], g["median_abs_rel_err"] * 100.0, marker="o", linewidth=2.0, label=friendly_method(method))
    ax.set_xscale("log")
    ax.set_xlabel("GSD (m)")
    ax.set_ylabel("Medyan mutlak goreli hata (%)")
    ax.set_title("Tum batch icin medyan hata - GSD iliskisi")
    ax.legend()
    fig.tight_layout()
    fig.savefig(error_chart, dpi=180)
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(8.8, 5.2))
    for method, g in summary.groupby("method", sort=False):
        ax.plot(g["gsd_m"], g["median_runtime_sec"], marker="o", linewidth=2.0, label=friendly_method(method))
    ax.set_xscale("log")
    ax.set_yscale("log")
    ax.set_xlabel("GSD (m)")
    ax.set_ylabel("Medyan runtime (sn)")
    ax.set_title("Tum batch icin medyan runtime - GSD iliskisi")
    ax.legend()
    fig.tight_layout()
    fig.savefig(runtime_chart, dpi=180)
    plt.close(fig)

    native = summary.loc[summary["gsd_m"] == summary["gsd_m"].min()].copy()
    fig, ax = plt.subplots(figsize=(8.3, 5.0))
    ax.scatter(native["median_runtime_sec"], native["median_abs_rel_err"] * 100.0, s=115, color="#1f6fb2")
    for _, row in native.iterrows():
        ax.annotate(
            friendly_method(row["method"]),
            (row["median_runtime_sec"], row["median_abs_rel_err"] * 100.0),
            textcoords="offset points",
            xytext=(8, 6),
            fontsize=9,
        )
    ax.set_xscale("log")
    ax.set_yscale("log")
    ax.set_xlabel("Medyan runtime (sn)")
    ax.set_ylabel("Medyan mutlak goreli hata (%)")
    ax.set_title("Native GSD icin dogruluk - hiz trade-off'u")
    fig.tight_layout()
    fig.savefig(tradeoff_chart, dpi=180)
    plt.close(fig)

    families = ["Analitik", "Gercekci", "Test pattern"]
    methods = list(native_family_summary["method"].unique())
    pivot = (
        native_family_summary.pivot(index="family", columns="method", values="median_abs_err")
        .reindex(families)
        .reindex(columns=methods)
        * 100.0
    )

    fig, ax = plt.subplots(figsize=(9.2, 5.2))
    x = range(len(families))
    width = 0.18
    offsets = [(-1.5 + i) * width for i in range(len(methods))]
    colors = ["#1f6fb2", "#2f9e44", "#d9480f", "#7b2cbf"]
    for offset, color, method in zip(offsets, colors, methods):
        ax.bar([i + offset for i in x], pivot[method].tolist(), width=width, label=friendly_method(method), color=color)
    ax.set_xticks(list(x))
    ax.set_xticklabels(families)
    ax.set_ylabel("Native medyan mutlak goreli hata (%)")
    ax.set_title("Yuzey ailesine gore native hata dagilimi")
    ax.legend(fontsize=8)
    fig.tight_layout()
    fig.savefig(family_chart, dpi=180)
    plt.close(fig)

    return {
        "error": error_chart,
        "runtime": runtime_chart,
        "tradeoff": tradeoff_chart,
        "family": family_chart,
    }


def build_report(
    results_dir: Path,
    output_path: Path,
    batch_results_long: pd.DataFrame,
    batch_summary: pd.DataFrame,
    batch_runs: pd.DataFrame,
    run_infos: list[dict[str, object]],
    charts: dict[str, Path],
    stats: ReportStats,
) -> None:
    doc = Document()
    style_document(doc)
    add_title_page(doc, output_path.name)

    nonref = batch_results_long.loc[~batch_results_long["is_reference_row"].fillna(False)].copy()
    nonref["preset"] = nonref["area_id"].map(extract_preset)
    nonref["family"] = nonref["preset"].map(family_for_preset)
    nonref["abs_rel_err"] = nonref["comparison_abs_rel_err"]

    native_gsd = float(nonref["gsd_m"].min())
    native_rows = nonref.loc[nonref["gsd_m"] == native_gsd].copy()

    native_summary = batch_summary.loc[batch_summary["gsd_m"] == native_gsd].sort_values("median_abs_rel_err").copy()
    native_summary["runtime_ratio_vs_gradient"] = (
        native_summary["median_runtime_sec"]
        / float(native_summary.loc[native_summary["method"] == "gradient_multiplier", "median_runtime_sec"].iloc[0])
    )

    native_family_summary = (
        native_rows.groupby(["family", "method"], as_index=False)
        .agg(median_abs_err=("abs_rel_err", "median"), mean_abs_err=("abs_rel_err", "mean"))
    )

    best_native = native_rows.sort_values(["area_id", "abs_rel_err", "runtime_sec"]).groupby("area_id").first().reset_index()
    best_native_counts = best_native["method"].value_counts()
    worst_cases = native_rows.sort_values("abs_rel_err", ascending=False).loc[:, ["area_id", "method", "abs_rel_err", "runtime_sec"]].head(8).copy()

    methods_current = [friendly_method(m) for m in run_infos[0]["params"]["methods"]]
    versions = run_infos[0]["versions"]
    params = run_infos[0]["params"]

    doc.add_heading("1. Yonetici Ozeti", level=1)
    doc.add_paragraph(
        "Bu proje, DEM/DSM GeoTIFF verilerinden 3B yuzey alani (A3D) hesaplamaya odaklanan bir Python araci ve "
        "deney altyapisina donusmus durumda. Kod tabani; blok bazli raster isleme, sentetik benchmark uretimi, "
        "analitik ground truth, coklu GSD analizi, Excel/JSON/PNG ciktilari ve test kapsami gibi arastirma odakli "
        "bilesenleri bir arada sunuyor."
    )
    doc.add_paragraph(
        f"Mevcut ana batch ciktisi {stats.run_count} farkli yuzey senaryosunu kapsiyor: {stats.analytic_count} analitik "
        f"benchmark, {stats.realistic_count} gercekci arazi ve {stats.test_pattern_count} test pattern. Bu set "
        f"{stats.gsd_count} farkli GSD duzeyinde ve {stats.method_count} yontemle calistirilmis; boylece "
        f"{stats.non_reference_rows} yontem-sonuc satiri ve {stats.reference_rows} referans satiri uretilmis."
    )
    doc.add_paragraph(
        f"Ilk ana sonuc, native GSD ({native_gsd:g} m) seviyesinde {friendly_method('jenness_window_8tri')} yonteminin "
        f"en dusuk medyan mutlak goreli hata ile one cikmasi. Bu yontem 24 senaryonun "
        f"{int(best_native_counts.get('jenness_window_8tri', 0))} tanesinde en iyi sonuca ulasiyor. "
        f"{friendly_method('gradient_multiplier')} hiz tarafinda acik ara lider; buna karsin "
        f"{friendly_method('sector_adaptive_jenness_integral')} mevcut batch'te hesaplama maliyetini ciddi bicimde "
        "artirmasina ragmen ayni olcude ek kazanc uretmiyor."
    )
    doc.add_paragraph(
        "Cozunurluk buyudukce yontemler arasi hata farki azaliyor; bu da orta ve kaba GSD seviyelerinde ana "
        "belirleyicinin algoritmadan cok yeniden ornekleme ve bilgi kaybi oldugunu gosteriyor. Bu bulgu, ileride "
        "yazilacak bildiride 'yontem etkisi' ile 'cozunurluk etkisi'ni ayri tartismanin gerekli olduguna isaret ediyor."
    )

    doc.add_heading("2. Calismanin Amaci ve Kapsami", level=1)
    doc.add_paragraph(
        "Calismanin temel amaci, DEM/DSM tabanli 3B yuzey alani hesaplamasinda farkli sayisal yaklasimlarin "
        "dogruluk-hiz davranisini kontrollu benchmarklar uzerinde karsilastirmak ve ileride gercek arazi "
        "uygulamalarina tasinabilecek bir deney catisi kurmaktir."
    )
    add_bullets(
        doc,
        [
            "DEM/DSM rasterlarindan A3D hesaplayabilen tekrar uretilebilir bir Python araci olusturuldu.",
            "Sentetik benchmark uretimi icin gercekci, test-pattern ve analitik yuzey aileleri tanimlandi.",
            "Yontemlerin farkli GSD seviyelerinde davranisini izleyecek toplu batch akisi kuruldu.",
            "Sonuclar Excel workbook, run metadata ve grafiklerle raporlanabilir hale getirildi.",
            "ROI, multiscale ve analitik ground truth gibi ileri seviyeli bilesenler kod tabanina eklendi.",
        ],
    )

    doc.add_heading("3. Tamamlanan Teknik Bilesenler", level=1)
    doc.add_paragraph(
        "Kod tabani artik yalnizca tek bir hesaplama fonksiyonundan ibaret degil; deney, benchmark, raporlama ve "
        "dogrulama katmanlariyla birlikte arastirma altyapisi karakteri tasiyor."
    )
    add_table(
        doc,
        headers=["Bilesen", "Rol"],
        rows=[[path, desc] for path, desc in MODULE_SUMMARY],
        col_widths_cm=[4.4, 11.4],
    )
    doc.add_paragraph()
    add_table(
        doc,
        headers=["Yontem", "Durum", "Kisa aciklama"],
        rows=[
            [
                friendly_method(method),
                "Bu batch'te aktif" if method in run_infos[0]["params"]["methods"] else "Kodda mevcut",
                METHOD_NOTES[method],
            ]
            for method in METHOD_LABELS
        ],
        col_widths_cm=[4.6, 3.0, 8.2],
    )

    doc.add_heading("4. Gelistirme Kilometre Taslari", level=1)
    doc.add_paragraph(
        "Git gecmisine gore ozellikle 2026-03-14 ile 2026-03-17 arasinda proje hizla olgunlasmis. "
        "Asagidaki maddeler, bu rapora temel olan kritik ilerlemeleri ozetliyor."
    )
    add_bullets(
        doc,
        [
            "2026-03-14: analitik yuzey destegi, synthetic reference kolonlari ve yeni CLI secenekleri eklendi.",
            "2026-03-14: plotting ve yeni referans metrikleri guclendirildi.",
            "2026-03-15: sonuc isleme akislari, DEM-list destegi ve Excel ciktilari gelistirildi.",
            "2026-03-17: Excel workbook bicimlendirmesi ve tablo destegi iyilestirildi.",
            "2026-03-17: allow_upsample ve GSD davranisi netlestirildi.",
            "2026-03-17: batch_summary.xlsx uretimi ve toplu hata ozetleri eklendi.",
            "2026-03-17: surface excess grafikleri ve iliskili cikti zenginlestirmeleri tamamlandi.",
        ],
    )

    doc.add_heading("5. Mevcut Benchmark Veri Seti", level=1)
    scenario_rows = [
        ["Analitik", str(stats.analytic_count), ", ".join(sorted(friendly_preset(p) for p in nonref["preset"].unique() if p.startswith("analytic_")))],
        ["Gercekci", str(stats.realistic_count), ", ".join(sorted(friendly_preset(p) for p in nonref["preset"].unique() if p in REALISTIC_PRESETS))],
        ["Test pattern", str(stats.test_pattern_count), ", ".join(sorted(friendly_preset(p) for p in nonref["preset"].unique() if p in TEST_PRESETS))],
    ]
    doc.add_paragraph(
        "Ana batch seti ayni geometriye sahip rasterlar uzerinde kurulmus durumda; bu durum yontem karsilastirmasinda "
        "kontrollu bir deney ortami sagliyor."
    )
    add_table(
        doc,
        headers=["Ozellik", "Deger"],
        rows=[
            ["Ana batch kosu sayisi", str(stats.run_count)],
            ["Raster boyutu", f"{stats.width} x {stats.height} hucre"],
            ["Native piksel boyutu", f"{stats.dx:g} m x {stats.dy:g} m"],
            ["Mekansal extent", f"{stats.extent_width_m:.1f} m x {stats.extent_height_m:.1f} m"],
            ["GSD listesi", ", ".join(str(g) for g in params["gsd_list"])],
            ["Aktif yontemler", ", ".join(methods_current)],
            ["Resampling", str(params["resampling"])],
            ["Slope method", str(params["slope_method"])],
            ["Worker sayisi", str(params["workers"])],
            ["Native-grid referansi olan yuzey", str(stats.native_ref_count)],
            ["Continuous ground truth olan yuzey", str(stats.continuous_gt_count)],
        ],
        col_widths_cm=[5.1, 10.7],
    )
    doc.add_paragraph()
    add_table(doc, headers=["Yuzey ailesi", "Adet", "Presetler"], rows=scenario_rows, col_widths_cm=[2.8, 1.4, 11.6])

    doc.add_heading("6. Deney Tasarimi ve Uretim Akisi", level=1)
    doc.add_paragraph(
        "Deney akisi, once sentetik GeoTIFF'lerin ve referans JSON dosyalarinin uretilmesi, ardindan bu rasterlarin "
        "coklu GSD seviyelerinde secili yontemlerle islenmesi ve sonuclarin toplu workbook'a konsolide edilmesi "
        "seklinde ilerliyor. Mevcut batch'te kullanilan ana konfigurasyon tum kosular icin homojen."
    )
    add_bullets(
        doc,
        [
            "Varsayilan batch yontem seti kullanildi: Jenness 8-Triangle, Sector-Adaptive Jenness, TIN 2-Triangle, Gradient Multiplier.",
            "GSD seviyeleri native, 0.1, 0.5, 1, 2, 5, 10, 20 ve 50 metre olarak tarandi.",
            "Tum kosullarda bilinear yeniden ornekleme ve Horn slope kerneli kullanildi.",
            "Analitik yuzeylerde surekli ground truth, tum sentetik rasterlarda native-grid referansi mevcut.",
        ],
    )
    doc.add_paragraph(
        f"Toplam hesaplama suresi, yalnizca yontemlerin kendi compute sureleri toplandiginda {format_duration(stats.total_compute_sec)} ediyor. "
        f"Tekil GSD-resampling adimlari ayristirildiginda benzersiz yeniden ornekleme maliyeti yaklasik "
        f"{format_duration(stats.total_unique_resample_sec)} seviyesinde."
    )

    doc.add_heading("7. Ana Bulgular", level=1)
    doc.add_paragraph(
        "Bu bolum, bildirinin sonuclar kismina cekirdek olabilecek sayisal gozlemleri ozetliyor. Tum degerler "
        "sonuclar_out_/batch_summary.xlsx uzerinden hesaplandi."
    )
    add_table(
        doc,
        headers=["Yontem", "Native medyan hata", "Native ort. hata", "Native medyan runtime", "Gradient'e gore hiz orani"],
        rows=[
            [
                friendly_method(row["method"]),
                format_pct(row["median_abs_rel_err"], 4),
                format_pct(row["mean_abs_rel_err"], 4),
                format_duration(float(row["median_runtime_sec"])),
                f"{float(row['runtime_ratio_vs_gradient']):.1f}x",
            ]
            for _, row in native_summary.iterrows()
        ],
        col_widths_cm=[4.3, 3.0, 3.0, 3.1, 3.0],
    )
    doc.add_paragraph()
    add_bullets(
        doc,
        [
            f"{friendly_method('jenness_window_8tri')}, native GSD'de en dusuk medyan mutlak goreli hata ile birinci durumda ve {int(best_native_counts.get('jenness_window_8tri', 0))}/{stats.run_count} yuzeyde en iyi sonucu veriyor.",
            f"{friendly_method('gradient_multiplier')} native seviyede en hizli yontem; medyan runtime'i {format_duration(float(native_summary.loc[native_summary['method'] == 'gradient_multiplier', 'median_runtime_sec'].iloc[0]))}.",
            f"{friendly_method('sector_adaptive_jenness_integral')} mevcut batch'te native seviyede {format_duration(float(native_summary.loc[native_summary['method'] == 'sector_adaptive_jenness_integral', 'median_runtime_sec'].iloc[0]))} medyan sure ile en maliyetli yontem.",
            "GSD buyudukce tum yontemlerin hata egrileri birbirine yaklasiyor; bu durum orta ve kaba cozunurluklerde ana hata kaynaginin yontem seciminden cok resampling/olcek kaybi oldugunu gosteriyor.",
            "En zor native ornekler canyon, terraced, patchwork ve mixed tiplerinde gozukuyor; yani keskin gecisler ve karma desenler hata icin kritik.",
        ],
    )

    doc.add_paragraph()
    doc.add_picture(str(charts["error"]), width=Cm(15.8))
    add_caption(doc, "Sekil 1. Tum batch icin medyan mutlak goreli hata - GSD iliskisi.")
    doc.add_picture(str(charts["runtime"]), width=Cm(15.8))
    add_caption(doc, "Sekil 2. Tum batch icin medyan runtime - GSD iliskisi.")
    doc.add_picture(str(charts["tradeoff"]), width=Cm(14.8))
    add_caption(doc, "Sekil 3. Native GSD seviyesinde dogruluk - hiz trade-off'u.")
    doc.add_picture(str(charts["family"]), width=Cm(15.8))
    add_caption(doc, "Sekil 4. Yuzey ailesine gore native hata davranisi.")

    doc.add_heading("8. Yuzey Ailesi Bazli Gozlemler", level=1)
    add_table(
        doc,
        headers=["Yuzey ailesi", "Yontem", "Native medyan hata", "Native ortalama hata"],
        rows=[
            [row["family"], friendly_method(row["method"]), format_pct(row["median_abs_err"], 4), format_pct(row["mean_abs_err"], 4)]
            for _, row in native_family_summary.sort_values(["family", "median_abs_err"]).iterrows()
        ],
        col_widths_cm=[3.0, 4.3, 4.2, 4.2],
    )
    doc.add_paragraph()
    add_bullets(
        doc,
        [
            "Analitik yuzeylerde tum yontemler native seviyede birbirine cok yakin; bu, temel formulasyonlarin puruzsuz yuzeylerde tutarli calistigini gosteriyor.",
            "Gercekci ve test-pattern gruplarinda klasik Jenness acik bicimde daha dusuk median hataya iniyor.",
            "Patchwork, terraced ve canyon gibi keskin/karisik yapili yuzeyler ayriklastirma tercihlerini daha hassas hale getiriyor.",
        ],
    )

    doc.add_heading("9. Cikti Envanteri", level=1)
    doc.add_paragraph(
        f"{results_dir.name} klasoru yalnizca tekil sonuc dosyalari degil, bildirinin ek materyalini de uretebilecek zengin bir dosya yapisi sunuyor."
    )
    add_table(
        doc,
        headers=["Artefakt", "Aciklama"],
        rows=[
            ["batch_summary.xlsx", "Tum batch icin uzun tablo, ozetler, pivot hata ve pivot runtime sayfalari."],
            ["<run>/results.xlsx", "Tek bir yuzey kosusu icin GSD x yontem sonuclari ve ilgili sheet'ler."],
            ["<run>/run_info.json", "Cevre surumleri, raster metadata'si ve parametre kaydi."],
            ["<run>/*.png", "A3D, ratio, runtime, error ve surface excess grafikleri."],
            ["old/ alt klasoru", "Eski batch kosularinin arsivlenmis versiyonlari."],
            ["report_assets/", "Bu rapor icin ek olarak uretilen toplu grafik dosyalari."],
        ],
        col_widths_cm=[4.5, 11.2],
    )
    doc.add_paragraph(
        "Repository kokunde yer alan dag_dsm.tif ve vadi_dsm.tif dosyalari, sentetik benchmark'tan sonraki gercek veri fazi icin hazir durumda. "
        "Mevcut batch_summary ise sentetik/analitik kosulari ozetliyor."
    )

    doc.add_heading("10. Dogrulama Durumu ve Sinirliliklar", level=1)
    doc.add_paragraph(
        "Kod tabaninda 8 test dosyasina dagilmis 61 test fonksiyonu bulunuyor. Ancak bu rapor hazirlanirken ayni shell oturumunda `pytest -q` komutu "
        "calistirildi ve test toplama asamasinda `rasterio` modulunun mevcut ortama kurulu olmamasi nedeniyle calisma durdu. Bu durum raporun ana "
        "sayisal bulgularini degistirmiyor; cunku sonuc klasorlerindeki run_info kayitlari benchmarklarin rasterio bulunan baska bir ortamda calistigini gosteriyor."
    )
    add_table(
        doc,
        headers=["Kaynak", "Durum"],
        rows=[
            ["Kayitli benchmark ortami", f"Python {versions['python'].split('|')[0].strip()}, rasterio {versions['rasterio']}, numpy {versions['numpy']}, pandas {versions['pandas']}, matplotlib {versions['matplotlib']}"],
            ["Bu rapor sirasinda test denemesi", "pytest toplama asamasinda rasterio eksikligi nedeniyle durdu"],
            ["Ana metodolojik sinirlilik", "Mevcut batch agirlikli olarak sentetik ve analitik yuzeylere dayaniyor"],
            ["Yontem kapsami sinirliligi", "Kodda 7 yontem var; mevcut default batch 4 yontemle sinirli"],
            ["Saha/gercek veri sinirliligi", "dag_dsm ve vadi_dsm icin henuz ayni duzeyde batch raporu uretilmemis"],
        ],
        col_widths_cm=[4.7, 11.0],
    )

    doc.add_heading("11. Bildiriye Gecis Icin Onerilen Sonraki Adimlar", level=1)
    add_bullets(
        doc,
        [
            "Mevcut 24-yuzey benchmark setini dondurup bir 'v1 benchmark protocol' olarak adlandirmak.",
            "Gercek veri fazina gecerek dag_dsm.tif ve vadi_dsm.tif icin ayni GSD/yontem akisini calistirmak.",
            "Kodda mevcut ama batch'te kapali olan bilinear, adaptive bilinear ve multiscale yontemlerini secmeli alt deney olarak eklemek.",
            "Birden fazla seed ile tekrarlayan sentetik kosular yapip istatistiksel guven araliklari cikarmak.",
            "Bildiriyi 'problem - related work - methods - benchmark design - results - discussion' omurgasiyla kurmak.",
            "Sonuc bolumunde 'native dogruluk', 'coarse GSD davranisi', 'zorlu yuzey tipleri' ve 'runtime trade-off' alt basliklarini kullanmak.",
        ],
    )

    doc.add_heading("12. Bildiri Icin Kullanilabilecek Cekirdek Sonuclar", level=1)
    add_bullets(
        doc,
        [
            f"Jenness 8-Triangle native seviyede ana kazanan: medyan hata {format_pct(float(native_summary.iloc[0]['median_abs_rel_err']), 4)}.",
            f"Gradient Multiplier hiz lideri: native medyan runtime {format_duration(float(native_summary.loc[native_summary['method'] == 'gradient_multiplier', 'median_runtime_sec'].iloc[0]))}.",
            f"Sector-Adaptive Jenness mevcut ayarlarda native seviyede Gradient'e gore yaklasik {float(native_summary.loc[native_summary['method'] == 'sector_adaptive_jenness_integral', 'runtime_ratio_vs_gradient'].iloc[0]):.0f} kat daha yavas.",
            "GSD buyudukce hata egrileri neredeyse ust uste geliyor; bu, cozunurluk etkisinin yontem etkisini bastirdigi alanlar oldugunu gosteriyor.",
            "Canyon ve terraced gibi keskin gecisli yuzeyler yontem farklarini en belirgin hale getiren zorlu benchmarklar.",
        ],
    )

    doc.add_page_break()
    doc.add_heading("Ek A. Native Seviyede En Zor Ornekler", level=1)
    add_table(
        doc,
        headers=["Alan", "Yontem", "Mutlak goreli hata", "Runtime"],
        rows=[
            [friendly_preset(extract_preset(row["area_id"])), friendly_method(row["method"]), format_pct(row["abs_rel_err"], 4), format_duration(float(row["runtime_sec"]))]
            for _, row in worst_cases.iterrows()
        ],
        col_widths_cm=[4.5, 4.8, 3.8, 3.2],
    )

    doc.add_heading("Ek B. Raporun Tekrar Uretilebilirligi", level=1)
    doc.add_paragraph("Bu DOCX raporu tekrar uretilebilir bicimde hazirlandi. Asagidaki komut ayni klasor yapisi uzerinde raporu yeniden yazar:")
    code = doc.add_paragraph()
    run = code.add_run(f"python generate_progress_report.py --results-dir sonuclar_out_ --output {output_path.as_posix()}")
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)

    doc.save(output_path)


def main() -> None:
    args = parse_args()
    results_dir = Path(args.results_dir).resolve()
    if not results_dir.exists():
        raise FileNotFoundError(f"Results directory not found: {results_dir}")

    output_path = Path(args.output) if args.output else results_dir / "yuzey_alani_calisma_ara_raporu_2026-03-23.docx"
    output_path = output_path.resolve()

    batch_path = results_dir / "batch_summary.xlsx"
    if not batch_path.exists():
        raise FileNotFoundError(f"batch_summary.xlsx not found under: {results_dir}")

    batch_results_long = pd.read_excel(batch_path, sheet_name="batch_results_long")
    batch_summary = pd.read_excel(batch_path, sheet_name="batch_summary")
    batch_runs = pd.read_excel(batch_path, sheet_name="batch_runs")

    run_infos = load_run_infos(results_dir, batch_runs)
    if not run_infos:
        raise RuntimeError("No run_info.json files found for current batch runs.")

    stats = build_stats(batch_results_long, batch_runs, run_infos)

    nonref = batch_results_long.loc[~batch_results_long["is_reference_row"].fillna(False)].copy()
    nonref["preset"] = nonref["area_id"].map(extract_preset)
    nonref["family"] = nonref["preset"].map(family_for_preset)
    nonref["abs_rel_err"] = nonref["comparison_abs_rel_err"]
    native_gsd = float(nonref["gsd_m"].min())
    native_family_summary = (
        nonref.loc[nonref["gsd_m"] == native_gsd]
        .groupby(["family", "method"], as_index=False)
        .agg(median_abs_err=("abs_rel_err", "median"), mean_abs_err=("abs_rel_err", "mean"))
    )
    charts = generate_charts(results_dir / "report_assets", batch_summary, native_family_summary)

    build_report(
        results_dir=results_dir,
        output_path=output_path,
        batch_results_long=batch_results_long,
        batch_summary=batch_summary,
        batch_runs=batch_runs,
        run_infos=run_infos,
        charts=charts,
        stats=stats,
    )

    print(f"Report generated: {output_path}")


if __name__ == "__main__":
    main()
